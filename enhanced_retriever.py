
import os
import json
import pickle
import numpy as np
from typing import List, Dict, Tuple, Optional, Union
import logging
from dataclasses import dataclass
import requests
from datetime import datetime
import tempfile
import PyPDF2
import docx

# Core ML libraries
from sentence_transformers import SentenceTransformer
import faiss
from transformers import AutoTokenizer

# External API libraries
import wikipedia
from urllib.parse import quote

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class RetrievedDocument:
    """Data class for retrieved documents"""
    content: str
    source: str
    score: float
    metadata: Optional[Dict] = None

class BaseRetriever:
    """Base class for all retrievers"""

    def retrieve(self, query: str, top_k: int = 5) -> List[RetrievedDocument]:
        raise NotImplementedError("Subclasses must implement retrieve method")

class EnhancedVectorRetriever(BaseRetriever):
    """
    Enhanced FAISS-based vector retriever with document upload support
    """

    def __init__(self, 
                 model_name: str = "all-MiniLM-L6-v2",
                 index_path: str = "faiss_index.bin",
                 documents_path: str = "documents.pkl",
                 chunk_size: int = 512):

        self.model_name = model_name
        self.index_path = index_path
        self.documents_path = documents_path
        self.chunk_size = chunk_size

        # Initialize sentence transformer
        logger.info(f"Loading sentence transformer: {model_name}")
        self.encoder = SentenceTransformer(model_name)
        self.embedding_dim = self.encoder.get_sentence_embedding_dimension()

        # Initialize tokenizer for chunking
        self.tokenizer = AutoTokenizer.from_pretrained("distilbert-base-uncased")

        # FAISS index and documents store
        self.index = None
        self.documents = []

        # Try to load existing index
        self.load_index()

    def add_documents_from_files(self, file_paths: List[str]) -> int:
        """
        Add documents from uploaded files
        Returns number of documents added
        """
        texts = []
        metadatas = []

        for file_path in file_paths:
            try:
                content = self._extract_text_from_file(file_path)
                if content:
                    texts.append(content)
                    metadatas.append({
                        "source": os.path.basename(file_path),
                        "type": "uploaded_file",
                        "upload_date": datetime.now().isoformat(),
                        "file_path": file_path
                    })
                    logger.info(f"Extracted {len(content)} characters from {file_path}")
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                continue

        if texts:
            self._add_texts_to_index(texts, metadatas)
            logger.info(f"Added {len(texts)} documents to index")

        return len(texts)

    def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif ext == '.pdf':
            return self._extract_from_pdf(file_path)

        elif ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)

        else:
            logger.warning(f"Unsupported file format: {ext}")
            return ""

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ""

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""

    def _add_texts_to_index(self, texts: List[str], metadatas: List[Dict]):
        """Add new texts to existing index"""
        # Chunk the texts
        all_chunks = []
        chunk_metadatas = []

        for i, text in enumerate(texts):
            chunks = self.chunk_text(text)
            all_chunks.extend(chunks)

            base_metadata = metadatas[i]
            for j, chunk in enumerate(chunks):
                chunk_meta = base_metadata.copy()
                chunk_meta.update({"chunk_id": j, "chunk_text": chunk})
                chunk_metadatas.append(chunk_meta)

        if not all_chunks:
            return

        # Generate embeddings
        logger.info("Generating embeddings for new documents...")
        embeddings = self.encoder.encode(all_chunks, show_progress_bar=True)
        embeddings = np.array(embeddings).astype('float32')
        faiss.normalize_L2(embeddings)

        # Add to existing index or create new one
        if self.index is None:
            self.index = faiss.IndexFlatIP(self.embedding_dim)

        self.index.add(embeddings)

        # Add new documents
        new_docs = [
            RetrievedDocument(
                content=chunk,
                source="local_corpus",
                score=0.0,
                metadata=meta
            ) for chunk, meta in zip(all_chunks, chunk_metadatas)
        ]

        self.documents.extend(new_docs)

        # Save updated index
        self.save_index()

        logger.info(f"Added {len(all_chunks)} chunks to index (total: {len(self.documents)})")

    def chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces based on token count"""
        tokens = self.tokenizer.tokenize(text)
        chunks = []

        overlap = 50
        for i in range(0, len(tokens), self.chunk_size - overlap):
            chunk_tokens = tokens[i:i + self.chunk_size]
            chunk_text = self.tokenizer.convert_tokens_to_string(chunk_tokens)
            chunks.append(chunk_text)

        return chunks

    def build_index_from_texts(self, texts: List[str], metadatas: Optional[List[Dict]] = None):
        """Build FAISS index from a list of texts"""
        logger.info(f"Processing {len(texts)} documents...")

        all_chunks = []
        chunk_metadatas = []

        for i, text in enumerate(texts):
            chunks = self.chunk_text(text)
            all_chunks.extend(chunks)

            base_metadata = metadatas[i] if metadatas else {"doc_id": i}
            for j, chunk in enumerate(chunks):
                chunk_meta = base_metadata.copy()
                chunk_meta.update({"chunk_id": j, "chunk_text": chunk})
                chunk_metadatas.append(chunk_meta)

        logger.info(f"Created {len(all_chunks)} chunks")

        # Generate embeddings
        logger.info("Generating embeddings...")
        embeddings = self.encoder.encode(all_chunks, show_progress_bar=True)
        embeddings = np.array(embeddings).astype('float32')

        # Build FAISS index
        logger.info("Building FAISS index...")
        self.index = faiss.IndexFlatIP(self.embedding_dim)
        faiss.normalize_L2(embeddings)
        self.index.add(embeddings)

        # Store documents and metadata
        self.documents = [
            RetrievedDocument(
                content=chunk,
                source="local_corpus",
                score=0.0,
                metadata=meta
            ) for chunk, meta in zip(all_chunks, chunk_metadatas)
        ]

        logger.info(f"Built index with {self.index.ntotal} vectors")

    def save_index(self):
        """Save FAISS index and documents to disk"""
        if self.index is not None:
            faiss.write_index(self.index, self.index_path)
            with open(self.documents_path, 'wb') as f:
                pickle.dump(self.documents, f)
            logger.info(f"Saved index with {self.index.ntotal} vectors")

    def load_index(self):
        """Load FAISS index and documents from disk"""
        try:
            if os.path.exists(self.index_path) and os.path.exists(self.documents_path):
                self.index = faiss.read_index(self.index_path)
                with open(self.documents_path, 'rb') as f:
                    self.documents = pickle.load(f)
                logger.info(f"Loaded index with {self.index.ntotal} vectors")
                return True
        except Exception as e:
            logger.error(f"Error loading index: {e}")
        return False

    def retrieve(self, query: str, top_k: int = 5) -> List[RetrievedDocument]:
        """Retrieve top-k most similar documents for a query"""
        if self.index is None or len(self.documents) == 0:
            logger.warning("No documents in index")
            return []

        # Encode query
        query_embedding = self.encoder.encode([query])
        query_embedding = np.array(query_embedding).astype('float32')
        faiss.normalize_L2(query_embedding)

        # Search
        scores, indices = self.index.search(query_embedding, min(top_k, len(self.documents)))

        # Prepare results
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(self.documents) and score > 0:  # Valid index and positive score
                doc = self.documents[idx]
                result_doc = RetrievedDocument(
                    content=doc.content,
                    source=doc.source,
                    score=float(score),
                    metadata=doc.metadata
                )
                results.append(result_doc)

        return results

class EnhancedWikipediaRetriever(BaseRetriever):
    """
    Enhanced Wikipedia retriever with better current information handling
    """

    def __init__(self, lang: str = "en", max_results: int = 5):
        self.lang = lang
        self.max_results = max_results
        wikipedia.set_lang(lang)

    def retrieve(self, query: str, top_k: int = 3) -> List[RetrievedDocument]:
        """
        Retrieve documents from Wikipedia with enhanced current event handling
        """
        try:
            # Enhanced search terms for current events
            enhanced_query = self._enhance_query_for_current_events(query)

            # Search for pages
            search_results = wikipedia.search(enhanced_query, results=min(top_k, self.max_results))

            # If no results with enhanced query, try original
            if not search_results:
                search_results = wikipedia.search(query, results=min(top_k, self.max_results))

            documents = []
            for i, title in enumerate(search_results[:top_k]):
                try:
                    # Get more content for current topics
                    sentences = 5 if self._is_current_topic(query) else 3
                    summary = wikipedia.summary(title, sentences=sentences)

                    # Try to get the full page content for better context
                    try:
                        page = wikipedia.page(title)
                        full_content = f"{summary}\n\nAdditional Context:\n{page.content[:1000]}"
                    except:
                        full_content = summary

                    doc = RetrievedDocument(
                        content=full_content,
                        source=f"wikipedia:{title}",
                        score=1.0 - (i * 0.1),
                        metadata={
                            "title": title,
                            "source_type": "wikipedia",
                            "url": f"https://en.wikipedia.org/wiki/{quote(title)}",
                            "enhanced_search": enhanced_query != query,
                            "retrieval_date": datetime.now().isoformat()
                        }
                    )
                    documents.append(doc)

                except wikipedia.exceptions.DisambiguationError as e:
                    # Handle disambiguation by trying the first few options
                    for option in e.options[:2]:  # Try first 2 options
                        try:
                            summary = wikipedia.summary(option, sentences=3)
                            doc = RetrievedDocument(
                                content=summary,
                                source=f"wikipedia:{option}",
                                score=1.0 - (i * 0.1) - 0.05,  # Slightly lower score for disambiguation
                                metadata={
                                    "title": option,
                                    "source_type": "wikipedia",
                                    "url": f"https://en.wikipedia.org/wiki/{quote(option)}",
                                    "disambiguation": True
                                }
                            )
                            documents.append(doc)
                            break  # Just take the first successful one
                        except:
                            continue

                except wikipedia.exceptions.PageError:
                    continue

                except Exception as e:
                    logger.warning(f"Error retrieving Wikipedia page {title}: {e}")
                    continue

            return documents

        except Exception as e:
            logger.error(f"Wikipedia search error: {e}")
            return []

    def _enhance_query_for_current_events(self, query: str) -> str:
        """Enhance query for better current event retrieval"""
        query_lower = query.lower()

        # Political queries
        if any(term in query_lower for term in ['president', 'election', 'government', 'politics']):
            if 'united states' in query_lower or 'us president' in query_lower or 'american president' in query_lower:
                return f"{query} 2024 current"

        # Add current year for time-sensitive queries
        if any(term in query_lower for term in ['current', 'now', 'today', 'latest', 'recent']):
            return f"{query} 2024"

        return query

    def _is_current_topic(self, query: str) -> bool:
        """Check if query is about current/recent topics"""
        current_keywords = ['current', 'now', 'today', 'latest', 'recent', 'president', 'election', '2024', '2023']
        return any(keyword in query.lower() for keyword in current_keywords)
